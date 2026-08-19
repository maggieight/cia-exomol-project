from __future__ import annotations

import csv
import hashlib
import io
import json
import os
import shutil
import subprocess
import tempfile
from collections import defaultdict
from dataclasses import dataclass
from decimal import Decimal, InvalidOperation
from pathlib import Path, PurePosixPath
from typing import Any, Sequence

TEMPERATURES = (293, 350, 374, 402, 431, 472)
CSV_HEADER = (
    "temperature",
    "wavenumber",
    "self_continuum_cross_section",
    "uncertainty",
)
EXPECTED = {
    293: (295, 2010, 5610),
    350: (381, 2010, 9080),
    374: (527, 2010, 9080),
    402: (516, 2010, 9130),
    431: (543, 2010, 9590),
    472: (509, 2010, 9590),
}


class ExtractionError(ValueError):
    """Raised when the supplementary table violates its canonical structure."""


@dataclass(frozen=True)
class ExtractedRow:
    temperature: int
    wavenumber: str
    coefficient: str
    uncertainty: str


def sha256_file(path: Path) -> str:
    """Return a file SHA-256 digest."""
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def locate_soffice() -> Path:
    """Locate a usable LibreOffice command without relying on one platform path."""
    candidates = [
        shutil.which("soffice"),
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/Applications/LibreOfficeDev.app/Contents/MacOS/soffice",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).is_file():
            return Path(candidate)
    raise RuntimeError("LibreOffice soffice is required for binary DOC conversion")


def convert_doc_to_docx(source: Path, destination: Path, soffice: Path) -> str:
    """Convert a binary Word DOC to a temporary DOCX using an isolated profile."""
    profile = destination.parent / "libreoffice-profile"
    cache = destination.parent / "cache"
    profile.mkdir()
    cache.mkdir()
    version = subprocess.run(
        [str(soffice), "--version"], check=True, text=True, capture_output=True
    ).stdout.strip()
    environment = os.environ.copy()
    environment["XDG_CACHE_HOME"] = str(cache)
    result = subprocess.run(
        [
            str(soffice),
            f"-env:UserInstallation={profile.as_uri()}",
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(destination.parent),
            str(source),
        ],
        check=True,
        text=True,
        capture_output=True,
        env=environment,
    )
    generated = destination.parent / f"{source.stem}.docx"
    if not generated.is_file():
        raise RuntimeError(f"LibreOffice did not create DOCX: {result.stdout} {result.stderr}")
    if generated != destination:
        generated.rename(destination)
    return version


def extract_column_tokens(docx_path: Path) -> tuple[list[list[str]], dict[str, int]]:
    """Extract the 13 vertical token streams from the Word table data row."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - environment dependency
        raise RuntimeError("python-docx is required for Ptashnik extraction") from exc
    document = Document(docx_path)
    if len(document.tables) != 1:
        raise ExtractionError(f"expected 1 Word table, found {len(document.tables)}")
    table = document.tables[0]
    if len(table.rows) != 2 or len(table.columns) != 13:
        raise ExtractionError(
            f"expected 2 physical rows and 13 columns, found {len(table.rows)} and {len(table.columns)}"
        )
    columns = [
        [paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip()]
        for cell in table.rows[1].cells
    ]
    validate_column_tokens(columns)
    return columns, {
        "word_table_count": len(document.tables),
        "physical_row_count": len(table.rows),
        "physical_column_count": len(table.columns),
    }


def validate_column_tokens(columns: Sequence[Sequence[str]]) -> None:
    """Require exactly 13 equally sized logical columns."""
    if len(columns) != 13:
        raise ExtractionError(f"expected 13 logical columns, found {len(columns)}")
    lengths = [len(column) for column in columns]
    if len(set(lengths)) != 1:
        raise ExtractionError(f"logical column token counts differ: {lengths}")
    if lengths[0] != 633:
        raise ExtractionError(f"expected 633 tokens per column, found {lengths[0]}")


def _finite_nonnegative(token: str, label: str, temperature: int, wavenumber: str) -> None:
    try:
        value = Decimal(token)
    except InvalidOperation as exc:
        raise ExtractionError(
            f"invalid {label} at {temperature} K, {wavenumber} cm^-1: {token!r}"
        ) from exc
    if not value.is_finite() or value < 0:
        raise ExtractionError(
            f"{label} must be finite and nonnegative at {temperature} K, {wavenumber} cm^-1"
        )


def logical_rows(columns: Sequence[Sequence[str]]) -> list[ExtractedRow]:
    """Reassemble vertical Word tokens into sorted long-format observations."""
    validate_column_tokens(columns)
    rows: list[ExtractedRow] = []
    wavenumbers = columns[0]
    for temperature_index, temperature in enumerate(TEMPERATURES):
        coefficient_column = columns[1 + temperature_index * 2]
        uncertainty_column = columns[2 + temperature_index * 2]
        for wavenumber, coefficient, uncertainty in zip(
            wavenumbers, coefficient_column, uncertainty_column
        ):
            coefficient_missing = coefficient == "--"
            uncertainty_missing = uncertainty == "--"
            if coefficient_missing and uncertainty_missing:
                continue
            if coefficient_missing != uncertainty_missing:
                raise ExtractionError(
                    f"incomplete pair at {temperature} K, {wavenumber} cm^-1: "
                    f"Cs={coefficient!r}, uncertainty={uncertainty!r}"
                )
            if not wavenumber.isdigit():
                raise ExtractionError(f"wavenumber is not an integer token: {wavenumber!r}")
            _finite_nonnegative(coefficient, "Cs", temperature, wavenumber)
            _finite_nonnegative(uncertainty, "uncertainty", temperature, wavenumber)
            rows.append(ExtractedRow(temperature, wavenumber, coefficient, uncertainty))
    return sorted(rows, key=lambda row: (row.temperature, int(row.wavenumber)))


def csv_bytes(rows: Sequence[ExtractedRow]) -> bytes:
    """Serialize canonical rows without changing scientific-notation tokens."""
    stream = io.StringIO(newline="")
    writer = csv.writer(stream, lineterminator="\n")
    writer.writerow(CSV_HEADER)
    writer.writerows(
        (row.temperature, row.wavenumber, row.coefficient, row.uncertainty)
        for row in rows
    )
    return stream.getvalue().encode("utf-8")


def validate_rows(rows: Sequence[ExtractedRow]) -> dict[str, Any]:
    """Validate counts, ordering, uniqueness, ranges, and numerical values."""
    keys = [(row.temperature, int(row.wavenumber)) for row in rows]
    duplicates = len(keys) - len(set(keys))
    grouped: dict[int, list[ExtractedRow]] = defaultdict(list)
    negative = 0
    nonfinite = 0
    for row in rows:
        grouped[row.temperature].append(row)
        for token in (row.coefficient, row.uncertainty):
            value = Decimal(token)
            negative += int(value < 0)
            nonfinite += int(not value.is_finite())
    if set(grouped) != set(TEMPERATURES):
        raise ExtractionError(f"unexpected temperatures: {sorted(grouped)}")
    per_temperature: dict[str, dict[str, int]] = {}
    for temperature in TEMPERATURES:
        values = grouped[temperature]
        waves = [int(row.wavenumber) for row in values]
        if any(later <= earlier for earlier, later in zip(waves, waves[1:])):
            raise ExtractionError(f"wavenumbers are not strictly increasing at {temperature} K")
        actual = (len(values), min(waves), max(waves))
        if actual != EXPECTED[temperature]:
            raise ExtractionError(
                f"unexpected statistics at {temperature} K: {actual}, expected {EXPECTED[temperature]}"
            )
        per_temperature[str(temperature)] = {
            "point_count": actual[0],
            "min_wavenumber": actual[1],
            "max_wavenumber": actual[2],
        }
    if len(rows) != 2771:
        raise ExtractionError(f"expected 2771 data rows, found {len(rows)}")
    if duplicates or negative or nonfinite:
        raise ExtractionError(
            f"invalid rows: duplicates={duplicates}, negative={negative}, nonfinite={nonfinite}"
        )
    return {
        "csv_data_row_count": len(rows),
        "per_temperature": per_temperature,
        "incomplete_pairs": 0,
        "duplicate_points": duplicates,
        "negative_values": negative,
        "nonfinite_values": nonfinite,
        "min_temperature": min(grouped),
        "max_temperature": max(grouped),
        "min_wavenumber": min(int(row.wavenumber) for row in rows),
        "max_wavenumber": max(int(row.wavenumber) for row in rows),
    }


def extradata_json(rows: Sequence[ExtractedRow], species: dict[str, Any]) -> dict[str, Any]:
    """Build the canonical extradata descriptor from validated CSV rows."""
    stats = validate_rows(rows)
    h2o = species["h2o"]
    component = {
        "formula": h2o["formula"],
        "slug": h2o["slug"],
        "cas_registry_number": h2o["cas_registry_number"],
    }
    return {
        "collision_pair": {
            "formula": "H2O-H2O",
            "slug": "h2o-h2o",
            "active_species_status": "no_unique_active_species",
            "active_species": None,
            "collider": None,
            "components": [component, dict(component)],
        },
        "dataset": {
            "id": "ptashnik2011",
            "recommendation_status": "supplementary",
            "collision_induced_absorption_xsecs": {
                "min_temperature": stats["min_temperature"],
                "max_temperature": stats["max_temperature"],
                "min_wavenumber": stats["min_wavenumber"],
                "max_wavenumber": stats["max_wavenumber"],
                "spectral_averaging_width": 20,
                "nominal_wavenumber_step": 10,
                "units": {
                    "temperature": "K",
                    "wavenumber": "cm^-1",
                    "spectral_averaging_width": "cm^-1",
                    "nominal_wavenumber_step": "cm^-1",
                },
                "column_schemas": {
                    "self_continuum_with_absolute_uncertainty": [
                        {"name": "wavenumber", "units": "cm^-1"},
                        {
                            "name": "self_continuum_cross_section",
                            "units": "cm^2 molecule^-1 atm^-1",
                        },
                        {
                            "name": "uncertainty",
                            "units": "cm^2 molecule^-1 atm^-1",
                            "uncertainty_type": "absolute",
                            "applies_to": "self_continuum_cross_section",
                        },
                    ]
                },
                "default_column_schema": "self_continuum_with_absolute_uncertainty",
                "data_file": "H2O-H2O.csv",
            },
            "sources": ["ptashnik2011"],
        },
    }


def validate_extradata_json(
    payload: dict[str, Any], rows: Sequence[ExtractedRow], metadata_dir: Path
) -> None:
    """Validate source, species, schema, path, and CSV-derived ranges."""
    sources = json.loads((metadata_dir / "sources.json").read_text(encoding="utf-8"))
    species = json.loads((metadata_dir / "species.json").read_text(encoding="utf-8"))
    dataset = payload["dataset"]
    if payload != extradata_json(rows, species):
        raise ExtractionError("extradata JSON differs from the canonical CSV-derived structure")
    source_key = dataset["sources"][0]
    source = sources.get(source_key)
    if not source or source.get("verified") is not True or source.get("ref") is not None:
        raise ExtractionError("ptashnik2011 source must exist, be verified, and have null ref")
    data_file = dataset["collision_induced_absorption_xsecs"]["data_file"]
    safe = PurePosixPath(data_file)
    if safe.is_absolute() or ".." in safe.parts or len(safe.parts) != 1:
        raise ExtractionError("data_file must be a safe local relative path")


def convert(
    source: Path, output_dir: Path, metadata_dir: Path
) -> dict[str, Any]:
    """Convert and validate Ptashnik Data Set S1 into canonical CSV and JSON."""
    source_digest = sha256_file(source)
    soffice = locate_soffice()
    with tempfile.TemporaryDirectory(prefix="ptashnik2011-") as temporary:
        docx_path = Path(temporary) / "water-data.docx"
        tool_version = convert_doc_to_docx(source, docx_path, soffice)
        columns, word_stats = extract_column_tokens(docx_path)
    rows = logical_rows(columns)
    stats = validate_rows(rows)
    serialized_csv = csv_bytes(rows)
    if serialized_csv != csv_bytes(logical_rows(columns)):
        raise ExtractionError("repeated conversion is not byte-for-byte deterministic")
    species = json.loads((metadata_dir / "species.json").read_text(encoding="utf-8"))
    payload = extradata_json(rows, species)
    validate_extradata_json(payload, rows, metadata_dir)
    serialized_json = (
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n"
    ).encode("utf-8")
    output_dir.mkdir(parents=True, exist_ok=True)
    csv_path = output_dir / "H2O-H2O.csv"
    json_path = output_dir / "H2O-H2O.json"
    csv_path.write_bytes(serialized_csv)
    json_path.write_bytes(serialized_json)
    report = {
        "source_path": "source_data/non_hitran/ptashnik2011/water-data.doc",
        "source_sha256": source_digest,
        "conversion_tool": "LibreOffice soffice",
        "conversion_tool_version": tool_version,
        **word_stats,
        "logical_token_count_per_column": [len(column) for column in columns],
        **stats,
        "csv_total_line_count": len(rows) + 1,
        "csv_sha256": sha256_file(csv_path),
        "json_sha256": sha256_file(json_path),
        "deterministic_repeat_conversion": True,
        "validation_result": "passed",
    }
    return report
